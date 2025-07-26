import gradio as gr
from docx import Document
from datetime import datetime

def generate_docs(name, email, phone, address, linkedin, github, summary, skills, experience, projects, education, certifications, company, job_title):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    resume_file = f"Resume_{timestamp}.docx"
    cover_file = f"CoverLetter_{timestamp}.docx"

    # Create Resume
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"{email} | {phone} | {address}")
    doc.add_paragraph(f"{linkedin} | {github}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Skills", level=1)
    for skill in skills.split(","):
        doc.add_paragraph(skill.strip(), style='List Bullet')

    doc.add_heading("Experience", level=1)
    doc.add_paragraph(experience)

    doc.add_heading("Projects", level=1)
    doc.add_paragraph(projects)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.add_heading("Certifications", level=1)
    for cert in certifications.split(","):
        doc.add_paragraph(cert.strip(), style='List Bullet')

    doc.save(resume_file)

    # Create Cover Letter
    doc2 = Document()
    doc2.add_heading("Cover Letter", 0)
    doc2.add_paragraph(f"Dear Hiring Manager at {company},\n")
    doc2.add_paragraph(
        f"I am writing to express my interest in the {job_title} position at {company}. "
        f"My name is {name}, and I bring relevant experience, skills, and passion to the role. "
        f"I have expertise in {skills} and experience in {experience}."
    )
    doc2.add_paragraph(
        "I am confident that I would be a valuable addition to your team and would welcome the opportunity to discuss further.\n"
        "Thank you for your time and consideration."
    )
    doc2.add_paragraph(f"Sincerely,\n{name}")
    doc2.save(cover_file)

    return f"Resume created: {resume_file}", f"Cover Letter created: {cover_file}", resume_file, cover_file

# Gradio UI
interface = gr.Interface(
    fn=generate_docs,
    inputs=[
        gr.Textbox(label="Full Name"),
        gr.Textbox(label="Email"),
        gr.Textbox(label="Phone"),
        gr.Textbox(label="Address"),
        gr.Textbox(label="LinkedIn URL"),
        gr.Textbox(label="GitHub URL"),
        gr.Textbox(label="Professional Summary"),
        gr.Textbox(label="Skills (comma-separated)"),
        gr.Textbox(label="Experience Summary"),
        gr.Textbox(label="Projects Description"),
        gr.Textbox(label="Education"),
        gr.Textbox(label="Certifications (comma-separated)"),
        gr.Textbox(label="Company Applying To"),
        gr.Textbox(label="Job Title"),
    ],
    outputs=[
        gr.Textbox(label="Resume Status"),
        gr.Textbox(label="Cover Letter Status"),
        gr.File(label="Download Resume (.docx)"),
        gr.File(label="Download Cover Letter (.docx)")
    ],
    title="📄 Offline Resume & Cover Letter Generator",
    description="Professional Resume & Cover Letter based on structured input — completely offline!"
)

interface.launch()
